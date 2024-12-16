from docx import Document
import uuid
import os

path = r"H:\develop\Sap-scripting\automations\documentation\templates"

class TemplateFile:

    def __init__(self, folder="default", title="template", count=1):
        self.folder = f"{path}\{folder}"  
        self.title = title
        self.count = count  

    def create_template(self):
        print("Creando template...")
        try:
            # Crear un nuevo documento de Word
            doc = Document()
            doc.add_paragraph(f"\n{self.title}\n")

            table = doc.add_table(rows=self.count, cols=2)
            table.style = 'Table Grid'  # Estilo opcional para la tabla
            table.autofit = False  # Desactiva el ajuste automático
            table.columns[0].width = 3886200  # Ajustar el ancho de la primera columna (en EMUs)
            table.columns[1].width = 3886200

            # Agregar el contenido al documento
            for i in range(self.count):
                # Llenar las celdas de la tabla
                row = table.rows[i]
                row.cells[1].text = f"#SCREEN{i + 1}#"  # Contenido en la segunda columna

            # Generar un nombre de archivo único
            unique_filename = f"{self.title}_{uuid.uuid4()}.docx"

            # Crear la carpeta si no existe
            if not os.path.exists(self.folder):
                os.makedirs(self.folder)

            # Guardar el archivo en la carpeta creada
            unique_template_path = os.path.join(self.folder, unique_filename)
            doc.save(unique_template_path)

            print("Template creado exitosamente.")

            return (self.folder, unique_filename)
        except Exception as e:
            print(f"Error al crear el template: {e}")

# Crear una instancia de la clase y llamar al método create_template
# creator = TemplateFile()
# creator.create_template()
